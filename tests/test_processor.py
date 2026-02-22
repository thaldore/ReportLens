"""
ReportProcessor Testleri.
DOCX tablo, başlık, metadata parse testleri.
"""
import tempfile
from pathlib import Path

import pytest
from docx import Document

from core.processor import ReportProcessor
from core.vector_store import VectorStore


class TestMetadataParsing:
    """Dosya adından metadata çıkarma testleri."""

    def test_full_metadata(self):
        meta = VectorStore.parse_metadata("Fen_2024_Oz_Degerlendirme_Raporu.md")
        assert meta["birim"] == "Fen"
        assert meta["yil"] == "2024"
        assert meta["tur"] == "Oz_Degerlendirme_Raporu"
        assert meta["dosya_adi"] == "Fen_2024_Oz_Degerlendirme_Raporu.md"

    def test_iibf_metadata(self):
        meta = VectorStore.parse_metadata("IIBF_2023_Eylem_Plani_Izleme_Raporu.md")
        assert meta["birim"] == "IIBF"
        assert meta["yil"] == "2023"
        assert meta["tur"] == "Eylem_Plani_Izleme_Raporu"

    def test_minimal_metadata(self):
        meta = VectorStore.parse_metadata("SomeName.md")
        assert meta["birim"] == "SomeName"
        assert "yil" not in meta or meta.get("yil") is None or meta["yil"] == ""


class TestHeadingLevel:
    """Heading seviye tespiti testleri."""

    def test_heading_1(self):
        assert ReportProcessor._get_heading_level("Heading 1") == 1

    def test_heading_3(self):
        assert ReportProcessor._get_heading_level("Heading 3") == 3

    def test_heading_no_number(self):
        assert ReportProcessor._get_heading_level("Heading") == 1

    def test_heading_max_6(self):
        assert ReportProcessor._get_heading_level("Heading 9") <= 6


class TestDocxProcessing:
    """DOCX işleme testleri."""

    def test_docx_with_table(self):
        """DOCX tablosu Markdown tablosuna dönüşmeli."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Test DOCX oluştur
            doc = Document()
            doc.add_heading("Test Başlık", level=1)
            doc.add_paragraph("Normal paragraf.")

            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = "Başlık1"
            table.cell(0, 1).text = "Başlık2"
            table.cell(0, 2).text = "Başlık3"
            table.cell(1, 0).text = "Veri1"
            table.cell(1, 1).text = "Veri2"
            table.cell(1, 2).text = "Veri3"

            docx_path = Path(tmpdir) / "test.docx"
            doc.save(str(docx_path))

            # İşle
            output_path = Path(tmpdir) / "test.md"
            processor = ReportProcessor()
            processor.process_docx(docx_path, output_path)

            # Kontrol et
            content = output_path.read_text(encoding="utf-8")
            assert "# Test Başlık" in content
            assert "Normal paragraf" in content
            assert "Başlık1" in content
            assert "Veri1" in content
            assert "|" in content  # Markdown tablo formatı

    def test_docx_with_heading(self):
        """DOCX başlıkları Markdown heading formatına dönüşmeli."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_heading("Birinci Seviye", level=1)
            doc.add_heading("İkinci Seviye", level=2)
            doc.add_paragraph("İçerik.")

            docx_path = Path(tmpdir) / "test.docx"
            doc.save(str(docx_path))

            output_path = Path(tmpdir) / "test.md"
            processor = ReportProcessor()
            processor.process_docx(docx_path, output_path)

            content = output_path.read_text(encoding="utf-8")
            assert "# Birinci Seviye" in content
            assert "## İkinci Seviye" in content
